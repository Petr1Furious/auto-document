import multiprocessing
import time
import docx
import easywebdav
from easywebdav import OperationFailed
from openpyxl import load_workbook
from requests.exceptions import ConnectionError
import yaml
import os
import string
from urllib.parse import unquote

easywebdav.basestring = str
easywebdav.client.basestring = str


def make_defaults():
    existing_config = []
    if os.path.exists('config.yml'):
        parsed = yaml.safe_load(open('config.yml', 'r', encoding="utf-8"))
        if parsed is not None:
            existing_config = list(parsed)

    desc = ['Почта, к которой привязан яндекс диск',

            'Пароль для приложения. Никому не отправляйте этот пароль или этот файл!\n'
            'Чтобы сгенерировать, зайдите на https://passport.yandex.ru/profile, раздел "Пароли и авторизация", '
            '"Включить пароли приложений" (или "Пароли приложений" если пароли приложений использовались ранее), '
            '"Создать новый пароль", "Файлы" и введите любое название.\n'
            'После этого скопируйте созданный пароль внутрь кавычек на следующей строке.',

            'Путь к таблице на яндекс диске (например, "/Проекты/Таблица с заявками.xlsx")',

            'Название листа для студентов',

            'Название листа для руководителей',

            'Путь к папке для документов от руководителей на яндекс диске (например, "/Проекты/Руководители/")',

            'Путь к папке для документов от студентов на яндекс диске (например, "/Проекты/Руководители/")',

            'Начальный ряд для обработки в листе для студентов',

            'Конечный ряд для обработки в листе для студентов или 0, если нужно идти до конца таблицы',

            'Начальный ряд для обработки в листе для руководителей',

            'Конечный ряд для обработки в листе для руководителей или 0, если нужно идти до конца таблицы',

            'Количество потоков для загрузки документов, 0 для отключения лимита',

            'Перезаписать существующие документы\n'
            'False - если на яндекс диске уже есть файл с совпадающим названием, он не будет перезаписан '
            '(загружается быстрее)\n'
            'True - созданный документ с совпадающим названием перезапишет тот, который уже есть на яндекс диске']
    names = ['username', 'password', 'table-path', 'leaders-sheet-name', 'students-sheet-name',
             'leaders-documents-path', 'students-documents-path', 'leaders-first-row', 'leaders-last-row',
             'students-first-row', 'students-last-row', 'threads-count', 'overwrite-existing-documents']
    values = ['', '', '/Таблица с заявками.xlsx',
              'ВЫГРУЗКАЗаявки от руководителей', 'ВЫГРУЗКА Инициативные заявки', '/', '/', 2, 0, 2, 0, 10, False]

    f = open('config.yml', 'a', encoding="utf-8")

    for i in range(len(desc)):
        if names[i] not in existing_config:
            if desc[i] != '':
                lines = desc[i].split('\n')
                for line in lines:
                    f.write('# ' + line + '\n')
            f.write(names[i] + ': ')
            if isinstance(values[i], str):
                f.write('\'' + values[i] + '\'')
            else:
                f.write(str(values[i]))
            f.write('\n\n')

    if names[0] not in existing_config:
        print('Введите имя пользователя и пароль в config.yml')
        exit(0)


def load_config():
    return yaml.safe_load(open('config.yml', 'r', encoding="utf-8"))


def get_person_name(s):
    words = s.split(' ')
    return words[0] + ' ' + [word[0].upper() + '' for word in words[1:]]


def get_project_name(s):
    exceptions = ['a', 'an', 'the', 'about', 'above', 'across', 'after', 'against', 'along', 'among', 'around',
                  'at', 'before', 'behind', 'between', 'beyond', 'but', 'by', 'concerning', 'despite', 'down',
                  'during', 'except', 'following', 'for', 'from', 'in', 'including', 'into', 'like', 'near',
                  'of', 'off', 'on', 'onto', 'out', 'over', 'past', 'plus', 'since', 'throughout', 'to', 'towards',
                  'under', 'until', 'up', 'upon', 'with', 'within', 'without']
    words = s.split(' ')
    res_words = []
    for word in words:
        if word.lower() not in exceptions:
            res_words.append(string.capwords(word))
        else:
            res_words.append(word.lower())
    return ' '.join(res_words)


def join(*to_join, sep=' '):
    return sep.join([mystr(s) for s in to_join if mystr(s) != ''])


def mystr(cell_text):
    return '' if cell_text is None else str(cell_text)


def get_file_name_leader(name, surname, patronymic, project_name):
    return ('Описание проекта_' + mystr(name).capitalize() + ' ' + mystr(surname)[0].upper() + '.' +
            (' ' + mystr(patronymic)[0].upper() + '.' if patronymic is not None and patronymic != '-' else '') +
            '_' + mystr(project_name) + '.docx').replace('/', '_').replace('\\', '_')


def get_file_name_student(name, project_name):
    words = name.split(' ')
    name = ''
    if len(words) >= 1:
        name = words[0].capitalize()
        if len(words) >= 2:
            name += ' ' + '. '.join([word[0].upper() for word in words[1:]]) + '.'
        name += '_'
    return ('Описание проекта_' + name + mystr(project_name) + '.docx').replace('/', '_').replace('\\', '_')


def make_document_leader(row):
    if row[3] is None or row[4] is None:
        return

    leader_doc = docx.Document('template_leader.docx')

    rows = leader_doc.tables[0].rows
    rows[0].cells[1].text = mystr(row[25])
    rows[1].cells[1].text = get_project_name(mystr(row[26]))
    rows[2].cells[1].text = mystr(row[27])
    rows[3].cells[1].text = join(row[28], row[29])
    rows[5].cells[1].text = \
        '1' if mystr(row[28]) == 'Индивидуальный' or mystr(row[29]) == 'Индивидуальный' else mystr(row[30])
    rows[6].cells[1].text = mystr(row[3]) + ' ' + mystr(row[4]) + ' ' + ('' if row[5] == '-' else mystr(row[5]))
    rows[7].cells[1].text = join(row[10], row[11], row[17], row[18], row[19])
    rows[8].cells[1].text = mystr(row[3]) + ' ' + mystr(row[4]) + ' ' + ('' if row[5] == '-' else mystr(row[5]))
    rows[9].cells[1].text = mystr(row[32])
    rows[10].cells[1].text = mystr(row[33])
    rows[11].cells[1].text = mystr(row[34])
    rows[12].cells[1].text = mystr(row[35])
    rows[13].cells[1].text = mystr(row[41])
    rows[15].cells[1].text = mystr(row[31])
    rows[16].cells[1].text = 'На email ' + mystr(row[6]) + '\n\n1 сентября 2022 по 15 октября 2022'
    rows[17].cells[1].text = join(row[36], row[37], row[39], sep='\n')

    leader_doc.save('leaders_docs' + os.sep + get_file_name_leader(row[3], row[4], row[5], row[25]))


def make_document_student(row):
    if row[3] is None or row[4] is None:
        return

    student_doc = docx.Document('template_student.docx')

    rows = student_doc.tables[0].rows
    rows[0].cells[1].text = mystr(row[22])
    rows[1].cells[1].text = get_project_name(mystr(row[23]))
    rows[2].cells[1].text = mystr(row[24])
    rows[3].cells[1].text = join(row[25], row[26])
    rows[4].cells[1].text = \
        '1' if mystr(row[25]) == 'Индивидуальный' or mystr(row[26]) == 'Индивидуальный' else mystr(row[27])
    rows[5].cells[1].text = mystr(row[3]) + ' ' + mystr(row[4]) + ' ' + ('' if row[5] == '-' else mystr(row[5]))
    rows[6].cells[1].text = mystr(row[9]) + ', ' + mystr(row[10])
    rows[7].cells[1].text = mystr(row[12]) + '\n' + join(join(row[16], row[18], row[20], sep=' '),
                                                         mystr(row[21]), sep=', ')
    rows[8].cells[1].text = mystr(row[30])
    rows[9].cells[1].text = mystr(row[31])
    rows[10].cells[1].text = mystr(row[32])
    rows[11].cells[1].text = mystr(row[33])
    rows[12].cells[1].text = mystr(row[34])
    rows[14].cells[1].text = mystr(row[29])
    rows[15].cells[1].text = 'На email ' + mystr(row[6]) + '\n\n1 сентября 2022 по 15 октября 2022'

    student_doc.save('students_docs' + os.sep + get_file_name_student(row[12], row[22]))


def make_documents(leaders_sheet, func, first_row, last_row):
    values_table = [[c.value for c in r] for r in leaders_sheet.rows]

    size = len(values_table) if last_row == -1 else last_row + 1
    total = size - first_row
    for i in range(first_row, size):
        func(values_table[i])
        done = i - first_row + 1
        if done % 10 == 0:
            print(str(done) + ' / ' + str(total))
    if total % 10 != 0:
        print(str(total) + ' / ' + str(total))


def upload(webdav, local_path, remote_path):
    webdav.upload(local_path, remote_path)


def clear_files():
    if os.path.exists('table.xlsx'):
        os.remove('table.xlsx')

    if os.path.exists('leaders_docs'):
        for existing in os.listdir('leaders_docs'):
            os.remove('leaders_docs' + os.sep + existing)

        os.removedirs("leaders_docs")

    if os.path.exists('students_docs'):
        for existing in os.listdir('students_docs'):
            os.remove('students_docs' + os.sep + existing)
        os.removedirs("students_docs")


def do_threads(threads, threads_count):
    started_set = set()
    old_completed_count = -1
    while True:
        completed_count = 0
        running_count = 0
        for i in range(len(threads)):
            if not threads[i].is_alive() and i in started_set:
                completed_count += 1
            if threads[i].is_alive():
                running_count += 1

        if old_completed_count != completed_count:
            print(str(completed_count) + ' / ' + str(len(threads)))
            old_completed_count = completed_count

        if completed_count == len(threads):
            break

        if running_count < threads_count or threads_count == 0:
            for i in range(len(threads)):
                if i not in started_set:
                    started_set.add(i)
                    threads[i].start()
                    running_count += 1
                    if running_count == threads_count:
                        break

        time.sleep(0.1)


def fix_folder_path(path):
    if path == '' or path[0] != '/':
        path = '/' + path
    if path[len(path) - 1] != '/':
        path = path + '/'
    return path


def try_upload():
    try:
        config = load_config()

        try:
            webdav = easywebdav.connect('webdav.yandex.ru', username=config['username'],
                                        password=config['password'], protocol='https')
            webdav.ls()
        except ConnectionError:
            print('Отсутствует подключение к интернету.')
            return False
        except OperationFailed:
            print('Неверный логин или пароль.')
            return False

        clear_files()

        try:
            table_path = config['table-path']
            if table_path == '' or table_path[0] != '/':
                table_path = '/' + table_path
            webdav.download(table_path, 'table.xlsx')
        except OperationFailed:
            print('По заданному пути таблица не найдена.')
            return False

        workbook = load_workbook('table.xlsx')
        leaders_sheet = workbook[config['leaders-sheet-name']]
        students_sheet = workbook[config['students-sheet-name']]

        os.makedirs('leaders_docs')
        os.makedirs('students_docs')
        print('Генерация документов от руководителей.')
        make_documents(leaders_sheet, make_document_leader,
                       config['leaders-first-row'] - 1, config['leaders-last-row'] - 1)
        print('Генерация документов от студентов.')
        make_documents(students_sheet, make_document_student,
                       config['students-first-row'] - 1, config['students-last-row'] - 1)

        leaders_path = fix_folder_path(config['leaders-documents-path'])
        students_path = fix_folder_path(config['students-documents-path'])
        try:
            leaders_documents_list = webdav.ls(leaders_path)
            students_documents_list = webdav.ls(students_path)
        except OperationFailed:
            print('Папка для документов не найдена.')
            return False

        leaders_documents_names = [os.path.splitext(os.path.basename(unquote(x.name)))[0]
                                   for x in leaders_documents_list if
                                   os.path.splitext(os.path.basename(unquote(x.name)))[1] == '.docx']
        students_documents_names = [os.path.splitext(os.path.basename(unquote(x.name)))[0]
                                    for x in students_documents_list if
                                    os.path.splitext(os.path.basename(unquote(x.name)))[1] == '.docx']

        threads = []
        for file in os.listdir('leaders_docs'):
            if config['overwrite-existing-documents'] or leaders_documents_names.count(os.path.splitext(file)[0]) == 0:
                threads.append(multiprocessing.Process(
                    target=upload,
                    args=(webdav, 'leaders_docs' + os.sep + file, leaders_path + file)))
        for file in os.listdir('students_docs'):
            if config['overwrite-existing-documents'] or students_documents_names.count(os.path.splitext(file)[0]) == 0:
                threads.append(multiprocessing.Process(
                    target=upload,
                    args=(webdav, 'students_docs' + os.sep + file, students_path + file)))

        print('Загрузка документов на диск.')
        do_threads(threads, config['threads-count'])

        os.remove('table.xlsx')
    except OperationFailed as e:
        print('Ошибка во время загрузки: ' + str(e))
        return False
    except ConnectionError as e:
        print('Ошибка подключения: ' + str(e))
        return False
    except Exception as e:
        print('Ошибка: ' + str(e))
        return False

    return True


if __name__ == '__main__':
    multiprocessing.freeze_support()
    make_defaults()

    start_time = time.time()
    if try_upload():
        print('Документы успешно созданы и загружены.')
    else:
        print('Документы не были созданы или загружены.')
    print('Время выполнения: ' + str(round(time.time() - start_time, 2)) + ' с.')

    clear_files()
